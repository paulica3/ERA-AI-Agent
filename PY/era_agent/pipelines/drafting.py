"""Contract drafting pipeline — fills ERA's DOCX template with client-specific data."""

import io
import re
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from era_agent.client import send_message

TEMPLATE_DIR = Path(__file__).resolve().parent.parent.parent / "templates"
CONTRACT_CLIENT_TEMPLATE = TEMPLATE_DIR / "contract_client.docx"

DRAFTING_SYSTEM = (
    "Ești un avocat expert în dreptul moldovenesc pentru firma Efrim Roșca & Asociații. "
    "Redactezi secțiunile specifice ale contractelor juridice cu terminologie juridică precisă "
    "și diacritice corecte. Răspunzi exclusiv în română."
)

MONTHS_RO = [
    "", "ianuarie", "februarie", "martie", "aprilie", "mai", "iunie",
    "iulie", "august", "septembrie", "octombrie", "noiembrie", "decembrie",
]


def draft_contract(
    client_name: str,
    client_type: str,
    client_idno: str,
    client_address: str,
    client_rep: str,
    client_rep_role: str,
    scope: str,
    services: str = "",
    fees: str = "",
    duration: str = "",
    contract_number: str = "",
) -> bytes:
    doc = Document(str(CONTRACT_CLIENT_TEMPLATE))

    today = date.today()
    date_str = f"{today.day} {MONTHS_RO[today.month]} {today.year}"
    num = contract_number.strip() or _auto_number()

    _replace_contract_header(doc, num, date_str)
    _replace_client_paragraph(doc, client_name, client_type, client_idno, client_address, client_rep, client_rep_role)
    _replace_scope_and_services(doc, scope, services, client_name)
    _replace_duration(doc, duration)
    _replace_fees(doc, fees)
    _replace_registration(doc, num, date_str)
    _replace_annexes_client(doc, client_name, client_type, client_rep)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Header (Para 5): Nr. 619/23/Data 10 iulie 2023 ───────────────────────────

def _replace_contract_header(doc, number, date_str):
    para = _find_para(doc, "Nr. ")
    if not para:
        return
    # runs: 'Nr. ', '619', '/23/Data ', '10', '', 'iulie', ' 2023'
    # Simplest: collapse to one run with new text
    _set_para_text(para, f"Nr. {number}/Data {date_str}")


# ── Client paragraph (Para 10) ────────────────────────────────────────────────

def _replace_client_paragraph(doc, name, ctype, idno, address, rep, role):
    para = _find_para(doc, "DRA Draexlmaier Automotive")
    if not para:
        return
    client_text = (
        f'"{name}" {ctype}, IDNO {idno}, adresa juridica {address}, '
        f'reprezentată corespunzător prin dl. {rep}, '
        f'care acționează în calitate de {role} ( Clientul ),'
    )
    _set_para_text(para, client_text)


# ── Scope + services (Paras 16–21) ───────────────────────────────────────────

def _replace_scope_and_services(doc, scope_input, services_input, client_name):
    scope_text, services_list = _generate_scope_and_services(scope_input, services_input, client_name)

    scope_para = _find_para(doc, "Domeniul specific")
    if scope_para:
        _set_para_text(scope_para, scope_text)

    # Paras 18–21 are the service lines; find them by known anchor text
    service_anchors = [
        "Analiza materialelor dosarului",
        "redactarea și/sau revizuirea",
        "asistență juridică și reprezentare în cadrul",
        "alte servicii de asistenta",
    ]
    service_paras = []
    for anchor in service_anchors:
        p = _find_para(doc, anchor)
        if p:
            service_paras.append(p)

    # Replace found service paragraphs with new services; blank any extras
    for i, para in enumerate(service_paras):
        if i < len(services_list):
            _set_para_text(para, services_list[i])
        else:
            _set_para_text(para, "")


def _generate_scope_and_services(scope_input, services_input, client_name):
    """Ask Claude to produce a scope paragraph and a services list."""
    prompt = (
        "Redactează secțiunea de domeniu de aplicare pentru un Contract de Client al firmei "
        "Efrim Roșca & Asociații, pe baza informațiilor de mai jos.\n\n"
        f"INFORMAȚII DESPRE CAZ:\n{scope_input}\n\n"
    )
    if services_input.strip():
        prompt += f"SERVICII SPECIFICE SOLICITATE:\n{services_input}\n\n"

    prompt += (
        "Returnează STRICT în formatul următor (două secțiuni, fără alt text):\n\n"
        "SCOP:\n"
        "<Un paragraf care începe cu: 'Domeniul specific de aplicare al prezentului Contract "
        "de Client constă în servicii de asistență juridică și reprezentare în favoarea "
        f"Clientului' — urmat de descrierea specifică a cazului.>\n\n"
        "SERVICII:\n"
        "<3–5 linii, fiecare descriind un serviciu juridic specific pe care ERA îl va presta. "
        "Fiecare linie se termină cu punct și virgulă.>"
    )

    raw = send_message(prompt, system=DRAFTING_SYSTEM, max_tokens=1500, use_web_search=False)

    scope_text = ""
    services_list = []

    scope_match = re.search(r"SCOP:\s*\n(.+?)(?:\n\nSERVICII:|\Z)", raw, re.DOTALL)
    if scope_match:
        scope_text = scope_match.group(1).strip()

    services_match = re.search(r"SERVICII:\s*\n(.+)", raw, re.DOTALL)
    if services_match:
        raw_services = services_match.group(1).strip()
        services_list = [
            line.lstrip("-•* ").strip()
            for line in raw_services.split("\n")
            if line.strip()
        ]

    if not scope_text:
        scope_text = (
            "Domeniul specific de aplicare al prezentului Contract de Client constă în "
            f"servicii de asistență juridică și reprezentare în favoarea Clientului în "
            f"legătură cu: {scope_input}"
        )
    if not services_list:
        services_list = [
            "Analiza materialelor și documentației prezentate de Client;",
            "Formularea recomandărilor juridice și a poziției juridice a Clientului;",
            "Redactarea și/sau revizuirea documentelor juridice necesare;",
            "Asistență juridică și reprezentare în cadrul procedurilor relevante;",
            "Alte servicii de asistență juridică, conform solicitărilor specifice ale Clientului, acceptate de ERA.",
        ]

    return scope_text, services_list


# ── Duration (Para 24) ────────────────────────────────────────────────────────

def _replace_duration(doc, duration_input):
    para = _find_para(doc, "Prezentul Contract de Client este încheiat valabil")
    if not para or not duration_input.strip():
        return
    duration_text = (
        f"Prezentul Contract de Client este încheiat valabil astăzi, data semnării lui de "
        f"către ambele părți, și va rămâne în vigoare {duration_input.strip()}."
    )
    _set_para_text(para, duration_text)


# ── Fees (Paras 34–35) ────────────────────────────────────────────────────────

def _replace_fees(doc, fees_input):
    if not fees_input.strip():
        return
    # Replace the specific fee lines [a] and [b] with user-provided text
    # Split user input by newlines; if only one block, put it in [a] slot
    fee_lines = [l.strip() for l in fees_input.strip().split("\n") if l.strip()]

    fee_anchors = [
        "[a] pentru analiza",
        "[b] Pentru serviciile de reprezentare",
    ]
    for i, anchor in enumerate(fee_anchors):
        para = _find_para(doc, anchor)
        if para and i < len(fee_lines):
            label = f"[{'abcdefgh'[i]}] "
            _set_para_text(para, label + fee_lines[i])
        elif para and i >= len(fee_lines):
            _set_para_text(para, "")


# ── Registration line (Para 111) ─────────────────────────────────────────────

def _replace_registration(doc, number, date_str):
    para = _find_para(doc, "Contractul de client a fost înregistrat sub nr.")
    if not para:
        return
    _set_para_text(
        para,
        f"Contractul de client a fost înregistrat sub nr. {number} din {date_str} "
        f"în Registrul contractelor cu clienții.",
    )


# ── Annexes: replace all occurrences of the old client name ──────────────────

def _replace_annexes_client(doc, client_name, client_type, client_rep):
    old_name = "DRA Draexlmaier Automotive"
    old_rep  = "Alexandru BORDEIANU"
    new_full = f"{client_name}" if client_type.lower() in ("persoana fizica", "persoană fizică") \
               else f'"{client_name}" {client_type}'

    for para in doc.paragraphs:
        full = para.text
        if old_name in full or old_rep in full:
            new_text = full.replace(old_name + '" SRL', new_full) \
                           .replace('"' + old_name + '"', f'"{client_name}"') \
                           .replace(old_name, client_name) \
                           .replace(old_rep, client_rep)
            if new_text != full:
                _set_para_text(para, new_text)


# ── Utilities ─────────────────────────────────────────────────────────────────

def _find_para(doc, search_text):
    for para in doc.paragraphs:
        if search_text in para.text:
            return para
    return None


def _set_para_text(para, new_text):
    """Replace all runs in a paragraph with a single run containing new_text."""
    if not para.runs:
        para.add_run(new_text)
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _auto_number():
    today = date.today()
    return f"{today.year % 100:02d}/{today.month:02d}"
