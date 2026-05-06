"""Invoice drafting pipeline — fills ERA's invoice_client.docx template."""

import copy
import io
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

TEMPLATE_PATH = Path(__file__).parent.parent.parent / "templates" / "invoice_client.docx"


# ── Low-level helpers ────────────────────────────────────────────────────────

def _para_text(para) -> str:
    return "".join(r.text for r in para.runs)


def _set_para_text(para, text: str) -> None:
    """Collapse all runs into the first one, preserving its formatting."""
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for r in para.runs[1:]:
        r.text = ""


def _set_cell_text(cell, text: str, bold: bool | None = None) -> None:
    """Set the text of the first paragraph in a table cell."""
    para = cell.paragraphs[0]
    if not para.runs:
        run = para.add_run(text)
        if bold is not None:
            run.bold = bold
        return
    para.runs[0].text = text
    if bold is not None:
        para.runs[0].bold = bold
    for r in para.runs[1:]:
        r.text = ""


def _find_para(paragraphs, needle: str):
    """Return first paragraph whose text contains needle, or None."""
    for p in paragraphs:
        if needle in _para_text(p):
            return p
    return None


# ── Amount helpers ───────────────────────────────────────────────────────────

def _parse_amount(s: str) -> float:
    """Parse '5,500.00' or '5500' → 5500.0."""
    try:
        return float(re.sub(r"[^\d.]", "", s))
    except ValueError:
        return 0.0


def _fmt(amount: float, currency: str = "EUR") -> str:
    return f"{currency} {amount:,.2f}"


def _parse_expenses(text: str) -> list[tuple[str, float]]:
    """
    Parse textarea input: one expense per line, format 'Description: amount'.
    Returns list of (description, float_amount).
    """
    items = []
    for line in (text or "").strip().splitlines():
        line = line.strip()
        if not line:
            continue
        if ":" in line:
            desc, _, amt_str = line.rpartition(":")
            items.append((desc.strip(), _parse_amount(amt_str.strip())))
        else:
            items.append((line, 0.0))
    return items


# ── Expense row builder ──────────────────────────────────────────────────────

def _rebuild_expense_row(row, expenses: list[tuple[str, float]], currency: str) -> None:
    """
    Rewrite the out-of-pocket row with a variable number of expense items.
    Left cell: "Out of pockets:\n1. Name\n2. Name …"
    Right cell: one amount per line, bold.
    """
    desc_cell = row.cells[0]
    amt_cell = row.cells[1]

    if not expenses:
        # Clear both cells
        for cell in (desc_cell, amt_cell):
            for para in cell.paragraphs:
                _set_para_text(para, "")
        return

    # ── Left cell ────────────────────────────────────────────────────
    # Use first paragraph as the header, then clone it for each item.
    header_para = desc_cell.paragraphs[0]
    _set_para_text(header_para, "Out of pockets:")

    # Remove any extra paragraphs beyond the first
    for extra in desc_cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    # Clone the header paragraph element for each expense item
    for i, (desc, _) in enumerate(expenses, start=1):
        new_el = copy.deepcopy(header_para._element)
        # Set text in all <w:t> inside the cloned element
        for t_el in new_el.findall(f".//{qn('w:t')}"):
            t_el.text = f"{i}. {desc}"
            break  # only the first <w:t>
        # Clear remaining <w:t> elements
        for t_el in list(new_el.findall(f".//{qn('w:t')}"))[1:]:
            t_el.text = ""
        desc_cell._element.append(new_el)

    # ── Right cell ────────────────────────────────────────────────────
    # Clear existing paragraphs, then write one amount line per expense.
    existing = amt_cell.paragraphs
    template_para = existing[0]
    _set_para_text(template_para, "")   # first line empty (padding)

    # Remove all but first
    for extra in existing[1:]:
        extra._element.getparent().remove(extra._element)

    for _, amt in expenses:
        new_el = copy.deepcopy(template_para._element)
        for t_el in new_el.findall(f".//{qn('w:t')}"):
            t_el.text = _fmt(amt, currency)
            break
        amt_cell._element.append(new_el)


# ── Main entry point ─────────────────────────────────────────────────────────

def draft_invoice(
    date: str,
    company_name: str,
    legal_address: str = "",
    client_iban: str = "",
    reg_no: str = "",
    vat_no: str = "",
    invoice_number: str = "",
    contract_ref: str = "",
    service_description: str = "",
    legal_fee: str = "0",
    currency: str = "EUR",
    expenses_text: str = "",
    partner_name: str = "Oleg Efrim",
    partner_title: str = "Partner",
    partner_email: str = "oleg.efrim@era.md",
) -> bytes:
    doc = Document(TEMPLATE_PATH)

    # ── Amounts ──────────────────────────────────────────────────────
    fee_val = _parse_amount(legal_fee)
    expenses = _parse_expenses(expenses_text)
    expense_total = sum(amt for _, amt in expenses)
    grand_total = fee_val + expense_total

    # ── Date ─────────────────────────────────────────────────────────
    for para in doc.paragraphs:
        if "17 April 2021" in _para_text(para):
            for run in para.runs:
                if "17 April 2021" in run.text:
                    run.text = run.text.replace("17 April 2021", date)
                    break
            break

    # ── Client info table (Table 0) ──────────────────────────────────
    t0 = doc.tables[0]
    _set_cell_text(t0.rows[0].cells[1], company_name, bold=True)
    _set_cell_text(t0.rows[1].cells[1], legal_address)
    _set_cell_text(t0.rows[2].cells[1], client_iban)
    _set_cell_text(t0.rows[3].cells[1], reg_no)
    _set_cell_text(t0.rows[4].cells[1], vat_no)

    # ── Invoice number ────────────────────────────────────────────────
    p = _find_para(doc.paragraphs, "Invoice no.")
    if p:
        _set_para_text(p, f"Invoice no. {invoice_number}")

    # ── Contract reference ────────────────────────────────────────────
    p = _find_para(doc.paragraphs, "based on the")
    if p:
        _set_para_text(p, f"(based on the {contract_ref})")

    # ── Services / billing table (Table 1) ───────────────────────────
    t1 = doc.tables[1]

    # Row 1 — main service fee
    desc_cell = t1.rows[1].cells[0]
    amt_cell = t1.rows[1].cells[1]

    _set_para_text(desc_cell.paragraphs[0], service_description)
    # Remove extra paragraphs in description cell
    for extra in desc_cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    # Find the bold paragraph in the amount cell and set fee
    for para in amt_cell.paragraphs:
        if para.runs and any(r.bold for r in para.runs):
            _set_para_text(para, _fmt(fee_val, currency))
            break

    # Row 2 — out-of-pocket expenses
    _rebuild_expense_row(t1.rows[2], expenses, currency)

    # Row 3 — grand total (red bold, right-aligned)
    total_amt_cell = t1.rows[3].cells[1]
    for para in total_amt_cell.paragraphs:
        if para.runs and any(r.text.strip() for r in para.runs):
            _set_para_text(para, _fmt(grand_total, currency))
            break

    # ── Signatory (Table 2 — floating) ───────────────────────────────
    t2 = doc.tables[2]
    sig_cell = t2.rows[0].cells[0]
    for para in sig_cell.paragraphs:
        txt = _para_text(para).strip()
        if "Oleg Efrim" in txt:
            _set_para_text(para, partner_name)
        elif txt == "Partner":
            _set_para_text(para, partner_title)
        elif "oleg.efrim@era.md" in txt:
            _set_para_text(para, partner_email)

    # ── Serialise ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
